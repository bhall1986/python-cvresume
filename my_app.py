from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


document = Document()

# Profile Picture
document.add_picture('me.png', width=Inches(2.0))

# Contact Information
name = input('What is your name? ')
speak(
    'Nice to virtually meet you, ' + name + ' Please enter your phone number and email, as we will be using it to build your resume. ')
phone_number = input('Enter your phone number here: ')
email = input('Enter your email address here: ')
document.add_paragraph(name + ' | ' + phone_number + ' | ' + email + ' | ')

# about me
document.add_heading('About Me')
document.add_paragraph(
    input('Tell me about yourself? ')
)

# Work Experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter Company: ')
start_date = input('From Date: ')
end_date = input('To Date: ')

p.add_run(company + ' ').bold = True
p.add_run(start_date + '-' + end_date + '\n').italic = True

experience_details = input('Describe your experience at ' + company + ' ')
p.add_run(experience_details)

# More Experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()
        company = input('Enter Company: ')
        start_date = input('From Date: ')
        end_date = input('To Date: ')

        p.add_run(company + ' ').bold = True
        p.add_run(start_date + '-' + end_date + '\n').italic = True

        experience_details = input('Describe your experience at ' + company + ' ')
        p.add_run(experience_details)

    else:
        break

# My Skills
document.add_heading('My Skills')
skill = input('Please list one of your skills: ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

# More Skills
while True:
    has_more_skills = input(
        'Do you have more skills you would like to list? Yes or No ')
    if has_more_skills.lower() == 'yes':
        skill = input('Please list one of your skills: ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'

    else:
        break

# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV Generated using PyCharm Python Coding Script designed by Brandon Hall'
document.save('CV.docx')
